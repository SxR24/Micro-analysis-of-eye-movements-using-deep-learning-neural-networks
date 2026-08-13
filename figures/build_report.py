#!/usr/bin/env python3
"""
build_report.py
======================================================================
Assembles the project report as a .docx to the BS7130 formatting spec:
A4, 12 pt, 1.5 line spacing, page numbers starting at the title page,
Harvard (Leicester) citations.

Text lives in report_text.py so that prose can be edited without touching
document mechanics. Figures are read from figures/output/.

    python figures/build_report.py

Sohil Ananth, MSc Bioinformatics & CS, University of Leicester
"""
import os
import sys

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(HERE)
FIGDIR = os.path.join(HERE, "output")
sys.path.insert(0, HERE)

from report_text import (TITLE, AUTHOR_BLOCK, ABSTRACT, SECTIONS,
                         REFERENCES, ACKNOWLEDGEMENTS, APPENDICES)


# ----------------------------------------------------------------------
def field(paragraph, instr):
    """Insert a Word field code (used for page numbers and the TOC)."""
    r = paragraph.add_run()
    fld = OxmlElement("w:fldChar"); fld.set(qn("w:fldCharType"), "begin")
    r._r.append(fld)
    r2 = paragraph.add_run()
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve")
    it.text = instr
    r2._r.append(it)
    r3 = paragraph.add_run()
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "separate")
    r3._r.append(f2)
    r4 = paragraph.add_run()
    f3 = OxmlElement("w:fldChar"); f3.set(qn("w:fldCharType"), "end")
    r4._r.append(f3)


def setup(doc):
    """A4, 2.5 cm margins, 12 pt Times New Roman, 1.5 spacing throughout."""
    for s in doc.sections:
        s.page_width, s.page_height = Cm(21.0), Cm(29.7)
        s.left_margin = s.right_margin = Cm(2.5)
        s.top_margin = s.bottom_margin = Cm(2.5)

    st = doc.styles["Normal"]
    st.font.name = "Times New Roman"
    st.font.size = Pt(12)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    pf = st.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(6)
    pf.space_before = Pt(0)

    for name, size, bold in (("Heading 1", 14, True),
                             ("Heading 2", 12, True),
                             ("Heading 3", 12, False)):
        h = doc.styles[name]
        h.font.name = "Times New Roman"
        h.font.size = Pt(size)
        # Word heading styles inherit their typeface from the document theme,
        # so setting font.name alone leaves them in the theme sans-serif.
        # The rFonts element has to be written for each script explicitly.
        rf = h.element.get_or_add_rPr().get_or_add_rFonts()
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rf.set(qn(attr), "Times New Roman")
        h.font.bold = bold
        h.font.italic = (name == "Heading 3")
        h.font.color.rgb = RGBColor(0, 0, 0)
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        h.paragraph_format.keep_with_next = True


def page_numbers(doc):
    """Footer page number, centred. Numbering starts at the title page."""
    for s in doc.sections:
        p = s.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        field(p, "PAGE")
        for r in p.runs:
            r.font.size = Pt(11)
            r.font.name = "Times New Roman"


def heading(doc, text, level):
    """Add a heading and force the typeface on the runs.

    Setting it on the style is not enough: Word resolves heading fonts through
    the document theme, and LibreOffice follows suit, so the style-level font
    is overridden and the headings render in the theme sans-serif.
    """
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = "Times New Roman"
        rf = r._element.get_or_add_rPr().get_or_add_rFonts()
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rf.set(qn(attr), "Times New Roman")
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.font.size = Pt(14 if level == 1 else 12)
        r.font.bold = (level <= 2)
        r.font.italic = (level == 3)
    return h


def para(doc, text, style=None, align=None, italic=False, size=None,
         space_after=None, first_line_indent=None):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.font.italic = italic
    if size:
        run.font.size = Pt(size)
    if align is not None:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    return p


def add_figure(doc, fname, number, caption, width_cm=15.5):
    path = os.path.join(FIGDIR, fname)
    if not os.path.exists(path):
        print("  WARNING: missing figure %s" % fname)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.add_run().add_picture(path, width=Cm(width_cm))

    c = doc.add_paragraph()
    c.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    c.paragraph_format.space_after = Pt(12)
    c.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r1 = c.add_run("Figure %d. " % number)
    r1.font.bold = True
    r1.font.size = Pt(11)
    r2 = c.add_run(caption)
    r2.font.size = Pt(11)


def add_table(doc, header, rows, number, caption, widths=None):
    c = doc.add_paragraph()
    c.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    c.paragraph_format.space_before = Pt(10)
    c.paragraph_format.space_after = Pt(4)
    c.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r1 = c.add_run("Table %d. " % number); r1.font.bold = True; r1.font.size = Pt(11)
    r2 = c.add_run(caption); r2.font.size = Pt(11)

    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Table Grid"
    for i, h in enumerate(header):
        cell = t.rows[0].cells[i]
        cell.text = ""
        pp = cell.paragraphs[0]
        pp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pp.paragraph_format.space_after = Pt(2)
        rr = pp.add_run(h); rr.font.bold = True; rr.font.size = Pt(10.5)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            pp = cells[i].paragraphs[0]
            pp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            pp.paragraph_format.space_after = Pt(2)
            rr = pp.add_run(str(v)); rr.font.size = Pt(10.5)
    if widths:
        for r_ in t.rows:
            for i, w in enumerate(widths):
                r_.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


# ----------------------------------------------------------------------
def build():
    doc = Document()
    setup(doc)
    page_numbers(doc)

    # ---------------- title page ----------------
    for _ in range(4):
        doc.add_paragraph()
    para(doc, TITLE, align=WD_ALIGN_PARAGRAPH.CENTER, size=16,
         space_after=24).runs[0].font.bold = True
    for line in AUTHOR_BLOCK:
        para(doc, line, align=WD_ALIGN_PARAGRAPH.CENTER, size=12, space_after=4)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ---------------- abstract ----------------
    heading(doc, "Abstract", 1)
    for p in ABSTRACT:
        para(doc, p, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ---------------- table of contents ----------------
    heading(doc, "Table of Contents", 1)
    p = doc.add_paragraph()
    field(p, r'TOC \o "1-3" \h \z \u')
    para(doc, "(In Word: right-click the field above and choose Update Field "
              "to populate the contents.)", italic=True, size=10.5)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ---------------- body ----------------
    for sec in SECTIONS:
        heading(doc, sec["heading"], 1)
        emit(doc, sec["blocks"])

    # ---------------- acknowledgements ----------------
    heading(doc, "Acknowledgements", 1)
    for p in ACKNOWLEDGEMENTS:
        para(doc, p, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # ---------------- references ----------------
    heading(doc, "References", 1)
    for r in REFERENCES:
        pp = doc.add_paragraph()
        pp.paragraph_format.left_indent = Cm(1.0)
        pp.paragraph_format.first_line_indent = Cm(-1.0)
        pp.paragraph_format.space_after = Pt(8)
        pp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pp.add_run(r)

    # ---------------- appendices ----------------
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    heading(doc, "Appendices", 1)
    emit(doc, APPENDICES)

    out = os.path.join(ROOT, "Report_Ananth_BS7130.docx")
    doc.save(out)
    print("Wrote", out)
    return out


def emit(doc, blocks):
    """Render a section body. Blocks are (kind, payload) tuples."""
    for kind, payload in blocks:
        if kind == "h2":
            heading(doc, payload, 2)
        elif kind == "h3":
            heading(doc, payload, 3)
        elif kind == "p":
            para(doc, payload, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        elif kind == "fig":
            add_figure(doc, *payload)
        elif kind == "table":
            add_table(doc, *payload)
        elif kind == "code":
            pp = doc.add_paragraph()
            pp.paragraph_format.left_indent = Cm(0.8)
            pp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            r = pp.add_run(payload)
            r.font.name = "Consolas"
            r.font.size = Pt(10)
        elif kind == "bullet":
            for b in payload:
                pp = doc.add_paragraph(style="List Bullet")
                pp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                pp.paragraph_format.space_after = Pt(3)
                pp.add_run(b)


if __name__ == "__main__":
    build()
